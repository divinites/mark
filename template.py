from docx import *
from docx.shared import Pt
from docx.shared import Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from matplotlib import font_manager
from libmark import *
from operator import itemgetter


class FontFeature:
    @staticmethod
    def system_font():
        raw_list = font_manager.findSystemFonts(fontpaths=None, fontext='ttf')
        font_list = []
        for font in raw_list:
            cutoff = font.rfind('/') + 1
            font_name = font[cutoff:-4]
            font_list.append(font_name)
        return font_list

    def __init__(self, *argv):
        font_list = self.system_font()
        self.alignment_dict = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
                               "center": WD_ALIGN_PARAGRAPH.CENTER, "just": WD_ALIGN_PARAGRAPH.JUSTIFY}
        self.flag = 0
        self.alignment = None
        self.treat = []
        self.pt = 14
        for arg in argv:
            if arg in font_list:
                self.font = arg
            else:
                self.font = 'Arial'
            if arg == 'b':
                self.treat.append('bold')
            if arg == 'u':
                self.treat.append('underline')
            if arg == 'i':
                self.treat.append('italic')

            if isinstance(arg, int):
                self.pt = arg

            if arg in self.alignment_dict.keys():
                self.alignment = self.alignment_dict[str(arg)]
            else:
                pass
            with open('tablestyle.list') as table_list:
                if arg in table_list:
                    self.table_style = arg
                    self.flag = 1
                else:
                    pass
        if self.flag == 0:
            self.table_style = "Light Shading"
        if self.alignment is None:
            self.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def font(self, font):
        self.font = font

    def pt(self, pt):
        self.pt = pt

    def alignment(self, alignment):
        self.alignment = self.alignment_dict[alignment]

    def table_style(self, table_style=None):
        if table_style is None:
            self.table_style = "Light Shading"
        self.table_style = table_style

    def add_treat(self, *treats):
        for i in treats:
            if i == 'b':
                self.treat.append('bold')
            if i == 'u':
                self.treat.append('underline')
            if i == 'i':
                self.treat.append('italic')

title_font = FontFeature(10, "center", 'b')
intro_font = FontFeature(10, "just")


class FeedbackForm:
    def __init__(self, file):

        self.document = Document()
        self.title = None
        self.result = Result(file)
        self.file = file

    def read_file(self, file):
        pass

    def add_title(self, title=None):
        if title is None:
            title = ''

        self.title = title

    def add_sentence(self, content=None, font_style=None, flag=None):
        if content is None:
            content = ''
        if font_style is None:
            font_style = intro_font
        if flag is None:
            flag = 'same'
        paragraph_list = self.document.paragraphs
        if not paragraph_list or flag == 'new':
            last_paragraph = self.document.add_paragraph()
        else:
            last_paragraph = paragraph_list.pop()
        last_paragraph.alignment = font_style.alignment
        sentence = last_paragraph.add_run(content)
        if len(font_style.treat) >= 1:
            for i in font_style.treat:
                exec("sentence.font."+i+"= True")
        sentence.font.name = font_style.font
        sentence.font.size = Pt(font_style.pt)

    def add_info(self, number, number_tag='Exam Number:', mark_tag='Mark:', font_style=None):
        if font_style is None:
            font_style = intro_font
        info = self.document.add_table(1, 4)
        info_cells = info.rows[0].cells
        s0 = info_cells[0].paragraphs[0].add_run(number_tag)
        info_cells[0].width = Emu(2500000)
        s0.font.size = Pt(font_style.pt)
        s1 = info_cells[1].paragraphs[0].add_run(number)
        info_cells[1].width = Emu(3400000)
        s1.font.size = Pt(font_style.pt)
        s1.font.underline = True
        s2 = info_cells[2].paragraphs[0].add_run(mark_tag)
        s2.font.size = Pt(font_style.pt)
        info_cells[2].width = Emu(360000)
        s3 = info_cells[3].paragraphs[0].add_run(str(self.result.round(number)))
        info_cells[3].width = Emu(360000)
        s3.font.size = Pt(font_style.pt)

    @staticmethod
    def print_stats(stat_table, cells, table_style=None):
        if table_style is None:
            table_style = FontFeature(10)
        for idx, stats in enumerate(stat_table):
            ptem = cells[idx].paragraphs[0]
            ptem.alignment = table_style.alignment
            tem = ptem.add_run(stats)
            tem.font.size = Pt(table_style.pt)
            tem.font.name = table_style.font
            tem.font.bold = False

    def add_form(self, std_no, table_style=None):
        if table_style is None:
            table_style = FontFeature(10)
        r_table_style = FontFeature(10, 'right')
        l_table_style = FontFeature(10, 'left')
        c_table_style = FontFeature(10, 'center')
        statements = file_process(self.file, 'f')
        pos_stats, std_cell, neg_stats = zip(*statements)
        form = self.document.add_table(len(statements), len(statements[0]))
        form.style = table_style.table_style
        pos_cells = form.columns[0].cells
        neg_cells = form.columns[2].cells
        mark_cells = form.columns[1].cells
        self.print_stats(neg_stats, neg_cells, l_table_style)
        self.print_stats(pos_stats, pos_cells, r_table_style)
        mark_stack = list(self.form_reformat())
        performance = self.result.get_mark(std_no)
        counter = 0
        for idx, value in enumerate(mark_stack):
            if value == '':
                mark_stack[idx] = self.draw_ticker(int(performance[counter]))
                counter += 1
            else:
                pass
        self.print_stats(mark_stack, mark_cells, c_table_style)

    def form_reformat(self):
        form_matrix = file_process(self.file, 'f')
        neg, mid, pos = zip(*form_matrix)
        return mid


    @staticmethod
    def draw_ticker(x, limit=5):
        vacant = u"☐ "
        ticker = u"☑ "
        ticker_list = []
        y = int(x)
        flag = y - 1
        for i in range(0, limit):
            if i == flag:
                ticker_list.append(ticker)
            else:
                ticker_list.append(vacant)
        return ' ' + ''.join(ticker_list) + ' '

    def new_page(self):
        self.document.add_page_break()

    def add_comment_table(self, std_no, category=None, table_style=FontFeature(10)):
        if category is None:
            category = (('Good Points:', 1, 5), ("Further Improvement", 2, 0),
                        ("Bonus Points", 3, 6), ("Additional Comment:", 3, ''))

        comments = self.document.add_table(len(category), 2)
        comments.style = table_style.table_style
        section_cells = comments.columns[0].cells
        content_cells = comments.columns[1].cells
        comment_stack = {}
        comment_type_counter = 0
        stud_info = self.result.get_student_info(std_no)

        for key, order, value in sorted(category, key=itemgetter(1)):
            comment_stack[key] = []
            for point, perf in stud_info["comment"]:
                if perf == value:
                    comment_stack[key].append(point)

            section_title = section_cells[comment_type_counter].paragraphs[0].add_run(key)
            section_title.font.size = Pt(table_style.pt)
            section_title.font.name = table_style.font
            for comment in comment_stack[key]:
                ps_comment = content_cells[comment_type_counter].add_paragraph(style='List Bullet').add_run(comment)
                ps_comment.bold = False
                ps_comment.font.size = Pt(10)
            comment_type_counter += 1


    def ug_form(self, stud_no, path=None):
        if path is None:
            path = './'

        self.add_title('Economics of Corporate Strategy - ASSIGNMENT FEEDBACK FORM 2014/15')
        self.add_sentence(self.title, title_font, 'new')
        self.add_sentence('This form is designed to provide you with specific '
                          'feedback on your own coursework assignment. The scale '
                          'below qualitatively presents an overview of the strengths'
                          ' and weaknesses of your work. Items are only ticked where'
                          ' applicable. Your tutor may provide additional comments overleaf.', intro_font, 'new')
        self.add_info(stud_no)
        self.add_form(stud_no)
        self.new_page()
        self.add_comment_table(stud_no)
        self.document.save(path + stud_no + '.docx')

    def print_form(self, stud_no, path, type):
        if type == "ecs":
            self.ug_form(stud_no, path)
        else:
            print("Now I only have ECS template ready...")






