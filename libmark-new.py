from docx import *
import csv
from docx.shared import Pt
from docx.shared import Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from xlrd import open_workbook
import operator


class student:
    def __init__(self, stdno):
        self.stdno = stdno


class feedback_form(Document):
    def __init__(self, title, intro):
        self.title = ''
        self.intro = ''

        
        
        
        
        
# Simply dealing with CSV
def DetectDelimiter(csvFile):
    with open(csvFile, 'r') as myCsvfile:
        header = myCsvfile.readline()
        if header.find(";") != -1:
            return ";"
        if header.find(",") != -1:
            return ","
    return ";"

def read_file(filenames):
    with open(filenames, 'rt') as csvfile:
        spamdata = csv.reader(csvfile.read().splitlines(), delimiter=DetectDelimiter(filenames))
    return spamdata

def order_file(target_file):
    target = read_file(target_file)
    target_matrix = []
    for i in target:
        target_matrix.append(i)
    return target_matrix

# Dealing with Excel workbook
def split_sheet(args):
    wb = open_workbook(args)
    mark_dict = {}
    for sheet in wb.sheets():
        if sheet.name == 'comment':
            mark_dict['comment'] = sheet
        if sheet.name == 'mark':
            mark_dict['mark'] = sheet
        if sheet.name == 'weight':
            mark_dict['weight'] = sheet
        if sheet.name == 'statement':
            mark_dict['statement'] = sheet
    return mark_dict


# Dealing with Excel sheets
def transfer_sheet(args):
    args_matrix = []
    for i in range(0, args.nrows):
        args_matrix.append(args.row_values(i))
    return args_matrix


def sort_grade(args):
    sorted_grade = sorted(args.items(), key=operator.itemgetter(1))
    return sorted_grade

# Draw Tickers
def draw_ticker(x):
    try:
        y = int(x)
        if y == 5:
            return u" ☑ ☐ ☐ ☐ ☐ "
        elif y == 4:
            return u" ☐ ☑ ☐ ☐ ☐ "
        elif y == 3:
            return u" ☐ ☐ ☑ ☐ ☐ "
        elif y == 2:
            return u" ☐ ☐ ☐ ☑ ☐ "
        elif y == 1:
            return u" ☐ ☐ ☐ ☐ ☑ "
        else:
            return y
    except:
        return x

""" Calculate the grade """
def grade(marks, weights):
    numbers = []
    for mark in marks:
        try:
            numbers.append(int(mark))
        except:
            pass

    grades =[]
    for i,j in zip(numbers, weights):
        grades.append(i*j)
    return int(round(sum(grades),0))

"""" Main function, process individual document"""

def doc_process(parameters):
    marks = parameters["mark"]
    comment_statement = parameters["all_comment"]
    individual_comment =parameters["comment"]
    weights_matrix = parameters["weight"]
    stat_matrix = parameters["statement"]

    additional_comment = individual_comment.pop()
    weights = []
    comment_weight = []

    for i in weights_matrix[0]:
        try:
            tmp = float(i)
            weights.append(tmp)
        except:
            pass
    for j in weights_matrix[1]:
        try:
            tmp = float(j)
            comment_weight.append(tmp)
        except:
            pass

    individual_comment = [int(i) for i in individual_comment]

    pos_stats, neg_stats = zip(*stat_matrix)
    document = Document()
    heads = document.add_paragraph()
    cheads = heads.add_run('Economics of Corporate Strategy - ASSIGNMENT FEEDBACK FORM 2014/15')
    cheads.bold = True
    heads.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cheads.font.name = 'Arial'
    cheads.font.size = Pt(14)

    intro = document.add_paragraph()

    intro_c = intro.add_run('This form is designed to provide you with specific '
                            'feedback on your own coursework assignment. The scale '
                            'below qualitatively presents an overview of the strengths'
                            ' and weaknesses of your work. Items are only ticked where'
                            ' applicable. Your tutor may provide additional comments overleaf.')

    std_no = marks.pop(0)


    intro_c.font.name = 'Arial'

    intro_c.font.size = Pt(14)

    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    document.add_paragraph()

    info = document.add_table(1, 4)
    info_cells = info.rows[0].cells
    S0 = info_cells[0].paragraphs[0].add_run('Examination Number:')
    info_cells[0].width = Emu(2500000)
    S0.font.size = Pt(14)

    S1 = info_cells[1].paragraphs[0].add_run(std_no)
    info_cells[1].width = Emu(3400000)
    S1.font.size = Pt(14)
    S1.font.underline = True
    S2 = info_cells[2].paragraphs[0].add_run('Mark:')
    S2.font.size = Pt(14)
    info_cells[2].width = Emu(360000)
    S3 = info_cells[3].paragraphs[0].add_run('')
    info_cells[3].width = Emu(360000)
    S3.font.size = Pt(14)
    S0.bold = True
    S1.bold = True
    S2.bold = True
    S3.bold = True
    results = document.add_table(19, 3)
    pos_cells = results.columns[0].cells
    neg_cells = results.columns[2].cells
    mark_cells = results.columns[1].cells
    for idx, stats in enumerate(neg_stats):
        ptem = neg_cells[idx].paragraphs[0]
        ptem.alignment = WD_ALIGN_PARAGRAPH.LEFT
        tem = ptem.add_run(stats)
        tem.font.size = Pt(10)
        tem.font.name = 'Arial'
        tem.bold = False
    for idx, stats in enumerate(pos_stats):
        ptem = pos_cells[idx].paragraphs[0]
        ptem.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        tem = ptem.add_run(stats)
        tem.font.size = Pt(10)
        tem.font.name = 'Arial'
        tem.bold = False

    for idx, mark in enumerate(marks):

        ptem = mark_cells[idx].paragraphs[0]
        ptem.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tem = ptem.add_run(draw_ticker(marks[idx]))
        tem.font.size = Pt(14)
        tem.font.name = 'Arial'
        tem.bold = False

    results.style = 'Light Shading'
    document.add_page_break()
    comment_section = document.add_paragraph().add_run('Comments:')
    comment_section.bold = True
    comment_section.font.size = Pt(14)
    comments = document.add_table(4, 2)
    comments.style = 'Light Shading'
    title_cells = comments.columns[0].cells
    comments_cells = comments.columns[1].cells
    title_size = title_cells[0].paragraphs[0].add_run('Good Points:')
    title_size.font.size = Pt(10)
    title_size = title_cells[1].paragraphs[0].add_run('Potential Improvements:')
    title_size.font.size = Pt(10)
    title_size = title_cells[3].paragraphs[0].add_run('Additional Comments:')
    title_size.font.size = Pt(10)

    for idx, i in enumerate(individual_comment):
        if i == 5:
            ps_comment = comments_cells[0].add_paragraph(style='List Bullet').add_run(comment_statement[idx])
            ps_comment.bold = False
            ps_comment.font.size =Pt(10)
        if i == 0:
            ng_comment = comments_cells[1].add_paragraph(style='List Bullet').add_run(comment_statement[idx])
            ng_comment.bold = False
            ng_comment.font.size = Pt(10)
        if i == 6:
            bs_comment = comments_cells[0].add_paragraph(style='List Bullet').add_run(comment_statement[idx]+' (Bonus Point)')
            bs_comment.bold = False
            bs_comment.font.size = Pt(10)
        else:
            pass

    ad_comment = comments_cells[3].add_paragraph(style='List Bullet').add_run(additional_comment)
    ad_comment.bold = False
    ad_comment.font.size = Pt(10)

    final_grade = grade(marks, weights)+ grade(individual_comment, comment_weight)
    while (final_grade % 10 >8):
        final_grade=final_grade+1
    finals = info_cells[3].paragraphs[0].add_run(str(final_grade))
    finals.font.size = Pt(14)
    finals.bold = True
    finals.font.underline = True

    document.save(std_no+'.docx')
    return std_no, final_grade