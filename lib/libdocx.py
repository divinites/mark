from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from matplotlib import font_manager


class FontFeature:
    @staticmethod
    def system_font():
        raw_list = font_manager.findSystemFonts(fontpaths=None, fontext='ttf')
        font_list = []
        for font in raw_list:
            cutoff = font.rfind('/') + 1
            font_name = font[cutoff:-4]
            font_list.append(font_name)
        return font_list

    def __init__(self, *argv):
        font_list = self.system_font()
        self.alignment_dict = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
                               "center": WD_ALIGN_PARAGRAPH.CENTER, "just": WD_ALIGN_PARAGRAPH.JUSTIFY}
        self.flag = 0
        self.alignment = None
        self.treat = []
        self.pt = 14
        for arg in argv:
            if arg in font_list:
                self.font = arg
            else:
                self.font = 'Arial'
            if arg == 'b':
                self.treat.append(('bold', True))
            if arg == 'u':
                self.treat.append(('underline', True))
            if arg == 'i':
                self.treat.append(('italic', True))
            if arg == 'nb':
                self.treat.append(('bold', False))

            if isinstance(arg, int):
                self.pt = arg

            if arg in self.alignment_dict.keys():
                self.alignment = self.alignment_dict[str(arg)]
            else:
                pass
            with open('tablestyle.list') as table_list:
                if arg in table_list:
                    self.table_style = arg
                    self.flag = 1
                else:
                    pass
        if self.flag == 0:
            self.table_style = "Light Shading"
        if self.alignment is None:
            self.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def font(self, font):
        self.font = font

    def pt(self, pt):
        self.pt = pt

    def alignment(self, alignment):
        self.alignment = self.alignment_dict[alignment]

    def table_style(self, table_style=None):
        if table_style is None:
            self.table_style = "Light Shading"
        self.table_style = table_style

    def add_treat(self, *treats):
        for i in treats:
            if i == 'b':
                self.treat.append(('bold', True))
            if i == 'u':
                self.treat.append(('underline', True))
            if i == 'i':
                self.treat.append(('italic', True))
            if i == 'nb':
                self.treat.append(('bold', False))


class DocWrapper:
    def __init__(self, document):
        self.default_font = FontFeature(10, 'nb')
        self.document = document

    def write(self, content=None, font_style=None, flag=None, last_paragraph=None):
        if content is None:
            content = ''
        if font_style is None:
            font_style = self.default_font
        if flag is None:
            flag = 'same'
        paragraph_list = self.document.paragraphs
        if last_paragraph is None:
            if not paragraph_list or flag == 'new':
                last_paragraph = self.document.add_paragraph()
            else:
                last_paragraph = paragraph_list.pop()
        last_paragraph.alignment = font_style.alignment
        sentence = last_paragraph.add_run(content)
        if len(font_style.treat) >= 1:
            for i, j in font_style.treat:
                exec("sentence.font." + i + "=" + str(j))
        sentence.font.name = font_style.font
        sentence.font.size = Pt(font_style.pt)

    def fill_cell(self, table, row, col, content, font_style=None):

        self.write(content, font_style, None, table.cell(row, col).paragraphs[0])

    def fill_row(self, table, row, content, font_style=None):

        if font_style is None:
            font_style = self.default_font
        for cell, item in zip(table.row_cells(row), content):
            self.write(item, font_style, None, cell.paragraphs[0])

    def fill_col(self, table, col, content, font_style=None):
        if font_style is None:
            font_style = self.default_font
        for cell, item in zip(table.column_cells(col), content):
            self.write(item, font_style, None, cell.paragraphs[0])

