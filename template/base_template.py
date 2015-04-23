from operator import itemgetter
from docx import Document
from lib.libdocx import *
from lib.libmark import *


title_font = FontFeature(10, "center", 'b')
intro_font = FontFeature(10, "just", 'nb')


class FeedbackForm:
    def __init__(self, file, stud_no):

        self.docx = None
        self.title = None
        self.result = Profiles(file)
        self.file = file
        self.stud_no = stud_no
        exec(stud_no + "= Document()")
        exec("self.docx=DocWrapper(" + stud_no + ")")
        assert isinstance(self.docx, DocWrapper)

    def read_file(self, file):
        pass

    def add_title(self, title=None):
        if title is None:
            title = ''

        self.title = title

    def add_info(self, number_tag='Exam Number:', mark_tag='Mark:', font_style=None):
        if font_style is None:
            font_style = intro_font
        row_content = [number_tag, self.stud_no, mark_tag, str(self.result.round(self.stud_no))]
        info = self.docx.document.add_table(1, 4)

        assert isinstance(info, object)
        self.docx.fill_row(info, 0, row_content, font_style)

    def add_form(self, table_style=None):
        if table_style is None:
            table_style = FontFeature(10)
        r_table_style = FontFeature(10, 'right', 'nb')
        l_table_style = FontFeature(10, 'left', 'nb')
        c_table_style = FontFeature(10, 'center')
        statements = file_process(self.file, 'f')
        pos_stats, std_cell, neg_stats = zip(*statements)
        form = self.docx.document.add_table(len(statements), len(statements[0]))
        form.style = table_style.table_style
        self.docx.fill_col(form, 2, neg_stats, l_table_style)
        self.docx.fill_col(form, 0, pos_stats, r_table_style)
        mark_stack = list(self.form_reformat())
        performance = self.result.get_mark(self.stud_no)
        counter = 0
        for idx, value in enumerate(mark_stack):
            if value == '':
                mark_stack[idx] = self.draw_ticker(int(performance[counter]))
                counter += 1
            else:
                pass
        self.docx.fill_col(form, 1, mark_stack, c_table_style)

    def form_reformat(self):
        form_matrix = file_process(self.file, 'f')
        neg, mid, pos = zip(*form_matrix)
        return mid


    @staticmethod
    def draw_ticker(x, limit=5):
        vacant = u"☐ "
        ticker = u"☑ "
        ticker_list = []
        y = int(x)
        flag = y - 1
        for i in range(0, limit):
            if i == flag:
                ticker_list.append(ticker)
            else:
                ticker_list.append(vacant)
        return ' ' + ''.join(ticker_list) + ' '

    def add_comment_table(self, category=None, table_style=FontFeature(10)):
        if category is None:
            category = (('Good Points:', 1, 5), ("Further Improvement", 2, 0),
                        ("Bonus Points", 3, 6), ("Additional Comment:", 3, ''))

        comments_table = self.docx.document.add_table(len(category), 2)
        comments_table.style = table_style.table_style
        section_cells = comments_table.columns[0].cells
        content_cells = comments_table.columns[1].cells
        comment_stack = {}
        comment_type_counter = 0
        stud_info = self.result.get_student_info(self.stud_no)

        for key, order, value in sorted(category, key=itemgetter(1)):
            comment_stack[key] = []
            for point, perf in stud_info["comment"]:
                if perf == value:
                    comment_stack[key].append(point)

            section_title = section_cells[comment_type_counter].paragraphs[0].add_run(key)
            section_title.font.size = Pt(table_style.pt)
            section_title.font.name = table_style.font
            for comment in comment_stack[key]:
                ps_comment = content_cells[comment_type_counter].add_paragraph(style='List Bullet').add_run(comment)
                ps_comment.bold = False
                ps_comment.font.size = Pt(10)
            comment_type_counter += 1

